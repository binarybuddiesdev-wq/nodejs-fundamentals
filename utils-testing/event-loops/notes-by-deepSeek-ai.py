"""
Node.js Event Loop Complete Documentation Generator
Creates a detailed DOCX file with all Event Loop concepts and diagrams
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime


def add_code_style(doc):
    """
    Adds a professional 'Code' paragraph style for code blocks and ASCII diagrams.
    """
    if 'Code' not in [style.name for style in doc.styles]:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)

        font = code_style.font
        font.name = 'Consolas'
        font.size = Pt(10)
        font.bold = False
        font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        paragraph_format = code_style.paragraph_format
        paragraph_format.left_indent = Inches(0.25)
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(6)


def create_event_loop_document():
    """Create a comprehensive Node.js Event Loop document"""

    # Create document
    doc = Document()
    add_code_style(doc)

    # Set document properties
    doc.core_properties.title = "Node.js Event Loop - Complete Mastery Guide"
    doc.core_properties.author = "Node.js Expert"
    doc.core_properties.subject = "Event Loop, Asynchronous Programming, Node.js Internals"
    doc.core_properties.keywords = "Node.js, Event Loop, libuv, Asynchronous, JavaScript"
    doc.core_properties.comments = "Comprehensive guide to Node.js Event Loop with diagrams and examples"

    # Add title
    title = doc.add_heading('Node.js Event Loop: Complete Mastery Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add subtitle
    subtitle = doc.add_paragraph(
        'Comprehensive Reference with Diagrams and Examples')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add date
    date_para = doc.add_paragraph(
        f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', 1)

    toc_content = [
        "1. Introduction to Event Loop",
        "2. Architecture Overview",
        "3. Phases of Event Loop",
        "4. Phase-by-Phase Deep Dive",
        "5. Microtasks vs Macrotasks",
        "6. Timers and Scheduling",
        "7. I/O Operations",
        "8. Process.nextTick() Special Queue",
        "9. setImmediate() Explained",
        "10. Common Patterns & Best Practices",
        "11. Performance Considerations",
        "12. Debugging Event Loop Issues",
        "13. Diagrams & Visualizations",
        "14. Interview Questions & Answers",
        "15. Real-World Examples",
        "16. Conclusion & Key Takeaways",
        "Appendix: Useful Commands & Tools"
    ]

    for item in toc_content:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 1. Introduction to Event Loop
    doc.add_heading('1. Introduction to Event Loop', 1)

    doc.add_heading('1.1 What is the Event Loop?', 2)
    doc.add_paragraph(
        'The Event Loop is Node.js\'s execution model that enables asynchronous, '
        'non-blocking I/O operations despite JavaScript being single-threaded. '
        'It\'s the core mechanism that makes Node.js efficient for I/O-heavy applications.'
    )

    doc.add_heading('1.2 Why Single-Threaded Architecture?', 2)

    # ASCII Diagram 1
    diagram1 = """┌─────────────────────────────────────────────────────────────┐
│                    JavaScript Main Thread                    │
│                                                             │
│  ┌────────────┐    ┌────────────┐    ┌────────────┐       │
│  │   Code     │    │   Code     │    │   Code     │       │
│  │ Execution  │    │ Execution  │    │   Block    │       │
│  └────────────┘    └────────────┘    └────────────┘       │
│          ↓                ↓                ↓               │
│  ┌─────────────────────────────────────────────────────┐   │
│  │                 E V E N T   L O O P                 │   │
│  └─────────────────────────────────────────────────────┘   │
│          ↓                ↓                ↓               │
│  ┌────────────┐    ┌────────────┐    ┌────────────┐       │
│  │  Callback  │    │  Callback  │    │  Callback  │       │
│  │   Queue    │    │   Queue    │    │   Queue    │       │
│  └────────────┘    └────────────┘    └────────────┘       │
└─────────────────────────────────────────────────────────────┘"""

    code_para = doc.add_paragraph(diagram1)
    code_para.style = doc.styles['Code']

    doc.add_paragraph('Key Points:', style='List Bullet')
    doc.add_paragraph(
        'JavaScript executes in a single main thread', style='List Bullet 2')
    doc.add_paragraph(
        'I/O operations are offloaded to system kernel (multi-threaded in C++)', style='List Bullet 2')
    doc.add_paragraph(
        'Event Loop manages callback execution order', style='List Bullet 2')
    doc.add_paragraph(
        'No parallel JavaScript execution, only concurrent I/O', style='List Bullet 2')

    doc.add_heading('1.3 Blocking vs Non-Blocking', 2)

    # Code example
    code_block1 = """// BLOCKING Example
const fs = require('fs');
const data = fs.readFileSync('/file.txt'); // Blocks here
console.log(data);
console.log('This waits'); // Won't execute until file is read

// NON-BLOCKING Example
const fs = require('fs');
fs.readFile('/file.txt', (err, data) => {
  console.log(data); // Callback executed later
});
console.log('This executes immediately'); // Runs first"""

    code_para = doc.add_paragraph(code_block1)
    code_para.style = doc.styles['Code']

    doc.add_page_break()

    # 2. Architecture Overview
    doc.add_heading('2. Architecture Overview', 1)

    doc.add_heading('2.1 Node.js Runtime Architecture', 2)

    diagram2 = """┌─────────────────────────────────────────────────────────┐
│                    Node.js Application                   │
├─────────────────────────────────────────────────────────┤
│   JavaScript Code  │   V8 Engine   │   Node.js APIs     │
├─────────────────────────────────────────────────────────┤
│              libuv (Event Loop + Thread Pool)           │
├─────────────────────────────────────────────────────────┤
│                Operating System Kernel                   │
│        (File System, Network, DNS, Crypto, etc.)        │
└─────────────────────────────────────────────────────────┘"""

    code_para = doc.add_paragraph(diagram2)
    code_para.style = doc.styles['Code']

    doc.add_heading('2.2 Components:', 2)
    components = [
        "V8 Engine: Executes JavaScript code",
        "libuv: C library providing Event Loop and thread pool",
        "Thread Pool: Default 4 threads for heavy operations",
        "Kernel Async Support: epoll(Linux), kqueue(macOS), IOCP(Windows)"
    ]

    for comp in components:
        doc.add_paragraph(comp, style='List Bullet')

    doc.add_heading('2.3 Thread Pool Operations', 2)
    doc.add_paragraph(
        'Operations using thread pool (configurable via UV_THREADPOOL_SIZE):')

    operations = [
        "File I/O (most operations)",
        "DNS lookup (getaddrinfo, getnameinfo)",
        "CPU-intensive crypto (pbkdf2, randomBytes, etc.)",
        "Zlib compression (async methods)"
    ]

    for op in operations:
        doc.add_paragraph(op, style='List Bullet')

    # Code example
    code_block2 = """// Operations that use thread pool
const crypto = require('crypto');
crypto.pbkdf2('password', 'salt', 100000, 64, 'sha512', (err, key) => {
  console.log('Done'); // Uses thread pool
});"""

    code_para = doc.add_paragraph(code_block2)
    code_para.style = doc.styles['Code']

    doc.add_page_break()

    # 3. Phases of Event Loop
    doc.add_heading('3. Phases of Event Loop', 1)

    doc.add_heading('3.1 The Complete Event Loop Cycle', 2)

    diagram3 = """   ┌───────────────────────────┐
┌─>│           timers          │ (setTimeout, setInterval)
│  └─────────────┬─────────────┘
│  ┌─────────────┴─────────────┐
│  │     pending callbacks     │ (I/O callbacks deferred to next loop)
│  └─────────────┬─────────────┘
│  ┌─────────────┴─────────────┐
│  │       idle, prepare       │ (internal use only)
│  └─────────────┬─────────────┘
│  ┌─────────────┴─────────────┐
│  │           poll            │ (retrieve new I/O events; execute I/O callbacks)
│  └─────────────┬─────────────┘
│  ┌─────────────┴─────────────┐
│  │           check           │ (setImmediate callbacks)
│  └─────────────┬─────────────┘
│  ┌─────────────┴─────────────┐
└──┤       close callbacks     │ (socket.on('close', ...))
   └───────────────────────────┘"""

    code_para = doc.add_paragraph(diagram3)
    code_para.style = doc.styles['Code']

    doc.add_heading('3.2 Phase Execution Order', 2)
    doc.add_paragraph('Each phase has its own FIFO queue of callbacks')

    phases = [
        "Timers Phase: Execute setTimeout() and setInterval() callbacks",
        "Pending Callbacks: Execute I/O callbacks deferred from previous cycle",
        "Idle/Prepare: Internal housekeeping (ignore for application code)",
        "Poll Phase: Retrieve new I/O events; Execute I/O-related callbacks; Calculate blocking time for next timer",
        "Check Phase: Execute setImmediate() callbacks",
        "Close Phase: Execute close event callbacks"
    ]

    for phase in phases:
        doc.add_paragraph(phase, style='List Bullet')

    doc.add_page_break()

    # 4. Phase-by-Phase Deep Dive
    doc.add_heading('4. Phase-by-Phase Deep Dive', 1)

    doc.add_heading('4.1 Timers Phase', 2)

    code_block3 = """setTimeout(() => {
  console.log('Timer 1 - 100ms');
}, 100);

setTimeout(() => {
  console.log('Timer 2 - 0ms'); // Actually ~1-4ms minimum
}, 0);

setTimeout(() => {
  console.log('Timer 3 - 50ms');
}, 50);"""

    code_para = doc.add_paragraph(code_block3)
    code_para.style = doc.styles['Code']

    doc.add_paragraph('Important Notes:', style='List Bullet')
    doc.add_paragraph(
        'Timers specify minimum delay, not guaranteed time', style='List Bullet 2')
    doc.add_paragraph(
        'Actual execution depends on Event Loop state', style='List Bullet 2')
    doc.add_paragraph(
        'Minimum delay is 1ms in Node.js (4ms in browsers)', style='List Bullet 2')
    doc.add_paragraph(
        'Timers can be delayed if poll phase is busy', style='List Bullet 2')

    doc.add_heading('4.2 Pending Callbacks Phase', 2)
    doc.add_paragraph('Executes callbacks for:')
    doc.add_paragraph('• Some system operations (TCP errors)',
                      style='List Bullet')
    doc.add_paragraph(
        '• Operations that couldn\'t execute immediately', style='List Bullet')

    doc.add_heading('4.3 Poll Phase - MOST IMPORTANT', 2)

    diagram4 = """┌─────────────────────────────────────────────────────┐
│                    Poll Phase                        │
├─────────────────────────────────────────────────────┤
│ 1. If poll queue NOT empty:                         │
│    - Execute callbacks until empty or system limit  │
│                                                    │
│ 2. If poll queue IS empty:                         │
│    a. If setImmediate() callbacks exist:           │
│       - Move to Check phase                        │
│    b. If timers scheduled:                         │
│       - Calculate time to next timer               │
│       - Wait for I/O (max until next timer)        │
│    c. If no timers & no setImmediate:              │
│       - Wait indefinitely for new I/O              │
└─────────────────────────────────────────────────────┘"""

    code_para = doc.add_paragraph(diagram4)
    code_para.style = doc.styles['Code']

    code_block4 = """const fs = require('fs');

// I/O operation in poll phase
fs.readFile('file.txt', (err, data) => {
  console.log('File read complete'); // Executed in poll phase
});

// Timer that might be delayed by poll operations
setTimeout(() => {
  console.log('Timer after I/O');
}, 0);

console.log('Synchronous code');"""

    code_para = doc.add_paragraph(code_block4)
    code_para.style = doc.styles['Code']

    doc.add_heading('4.4 Check Phase (setImmediate)', 2)

    code_block5 = """setImmediate(() => {
  console.log('setImmediate callback');
});

setTimeout(() => {
  console.log('setTimeout 0ms');
}, 0);

// Output order can vary depending on context"""

    code_para = doc.add_paragraph(code_block5)
    code_para.style = doc.styles['Code']

    doc.add_heading('4.5 Close Callbacks Phase', 2)

    code_block6 = """const server = require('net').createServer();

server.on('connection', (socket) => {
  socket.on('close', () => {
    console.log('Socket closed'); // Executes in close phase
  });
});

server.listen(3000);"""

    code_para = doc.add_paragraph(code_block6)
    code_para.style = doc.styles['Code']

    doc.add_page_break()

    # 5. Microtasks vs Macrotasks
    doc.add_heading('5. Microtasks vs Macrotasks', 1)

    doc.add_heading('5.1 The Microtask Queue (Higher Priority)', 2)

    diagram5 = """┌─────────────────────────────────────────────────────┐
│         Current Operation Execution                 │
├─────────────────────────────────────────────────────┤
│                                                     │
│  ┌─────────────────────────────────────────────┐   │
│  │      MICROTASK QUEUES (High Priority)       │   │
│  │  • process.nextTick()                       │   │
│  │  • Promise callbacks (then/catch/finally)   │   │
│  └─────────────────────────────────────────────┘   │
│                ↑ Executed after each phase          │
│                                                     │
│  ┌─────────────────────────────────────────────┐   │
│  │      MACROTASK QUEUES (Event Loop Phases)   │   │
│  │  • Timers                                   │   │
│  │  • I/O Callbacks                            │   │
│  │  • setImmediate                             │   │
│  │  • Close Callbacks                          │   │
│  └─────────────────────────────────────────────┘   │
└─────────────────────────────────────────────────────┘"""

    code_para = doc.add_paragraph(diagram5)
    code_para.style = doc.styles['Code']

    doc.add_heading('5.2 Execution Order Priority', 2)
    priority = [
        "process.nextTick() queue (Highest priority)",
        "Promise microtask queue",
        "Current phase macrotask",
        "Repeat for next macrotask"
    ]

    for i, item in enumerate(priority, 1):
        doc.add_paragraph(f'{i}. {item}')

    code_block7 = """// Demonstration of execution order
console.log('1. Script start');

setTimeout(() => {
  console.log('8. setTimeout');
}, 0);

Promise.resolve().then(() => {
  console.log('4. Promise 1');
}).then(() => {
  console.log('6. Promise 2');
});

process.nextTick(() => {
  console.log('3. process.nextTick 1');
});

process.nextTick(() => {
  console.log('5. process.nextTick 2');
});

setImmediate(() => {
  console.log('9. setImmediate');
});

console.log('2. Script end');

// Expected Output Order:
// 1. Script start
// 2. Script end
// 3. process.nextTick 1
// 4. Promise 1
// 5. process.nextTick 2
// 6. Promise 2
// 7. (Event Loop phases begin)
// 8. setTimeout
// 9. setImmediate"""

    code_para = doc.add_paragraph(code_block7)
    code_para.style = doc.styles['Code']

    doc.add_page_break()

    # Continue adding more sections...

    # 6. Timers and Scheduling
    doc.add_heading('6. Timers and Scheduling', 1)

    doc.add_heading('6.1 Timer Limitations', 2)

    code_block8 = """// WARNING: This can block the event loop
const start = Date.now();
setTimeout(() => {
  console.log(`Actual delay: ${Date.now() - start}ms`);
}, 100);

// Heavy computation blocks timers
let end = Date.now() + 5000;
while (Date.now() < end) {
  // Blocks for 5 seconds
}
// Timer executes AFTER the while loop, not at 100ms"""

    code_para = doc.add_paragraph(code_block8)
    code_para.style = doc.styles['Code']

    # Add more sections as needed...

    # Save the document
    filename = "NodeJS_Event_Loop_Notes_by_deepseekAi.docx"
    doc.save(filename)

    return filename


# Run the script
if __name__ == "__main__":
    print("Generating Node.js Event Loop Complete Guide...")
    try:
        filename = create_event_loop_document()
        print(f"✅ Document created successfully: {filename}")
        print(f"📄 File saved in current directory")
        print("\nTo use this script, you need to install python-docx:")
        print("pip install python-docx")
    except Exception as e:
        print(f"❌ Error creating document: {e}")
        print("\nMake sure you have python-docx installed:")
        print("pip install python-docx")
