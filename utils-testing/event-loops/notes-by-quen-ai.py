# nodejs_event_loop_notes.py

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text, style='Normal'):
    p = doc.add_paragraph(text, style=style)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.add_run(text).font.name = 'Consolas'
    p.runs[0].font.size = Pt(10)


doc = Document()
doc.core_properties.title = "Node.js Event Loop – Complete Technical Notes"
doc.core_properties.author = "Prepared for Mawa"

# Title
title = doc.add_heading('Node.js Event Loop – Complete Technical Notes', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("Prepared on: December 16, 2025", style='Intense Quote')
doc.add_page_break()

# Section 1
add_heading(doc, "1. What is the Event Loop?")
add_paragraph(doc, "The Event Loop is the core architectural feature that enables Node.js to handle asynchronous, non-blocking I/O operations efficiently—even though JavaScript in Node.js runs on a single main thread.")

add_paragraph(doc, "Node.js uses a combination of:")
add_paragraph(
    doc, "• V8 JavaScript Engine → Executes JavaScript code", style='List Bullet')
add_paragraph(
    doc, "• libuv → A C++ library that provides the event loop, thread pool, and async I/O handling", style='List Bullet')
add_paragraph(
    doc, "• C++ bindings and Node.js APIs → Bridge JavaScript with system-level operations", style='List Bullet')

doc.add_paragraph()

# Section 2
add_heading(doc, "2. Why is the Event Loop Needed?")
add_paragraph(doc, "JavaScript was originally designed to run in a single-threaded environment (e.g., browsers). In a server context like Node.js, blocking operations (e.g., file reads, network requests) would freeze the entire application if handled synchronously. The Event Loop solves this by:")
add_paragraph(
    doc, "• Delegating I/O operations to the OS kernel (which supports concurrency)", style='List Bullet')
add_paragraph(doc, "• Queueing callbacks to be executed later",
              style='List Bullet')
add_paragraph(
    doc, "• Continuously cycling through phases to process events", style='List Bullet')

doc.add_paragraph()

# Section 3
add_heading(doc, "3. Core Components of the Runtime")
add_paragraph(doc, "Three key components work together:")

add_paragraph(doc, "A) Call Stack", style='List Bullet')
add_paragraph(
    doc, "• LIFO structure that tracks function execution", style='List Bullet')
add_paragraph(doc, "• Synchronous code runs here directly",
              style='List Bullet')

add_paragraph(doc, "B) Node.js APIs (via libuv)", style='List Bullet')
add_paragraph(doc, "• Includes fs, net, dns, timers, etc.",
              style='List Bullet')
add_paragraph(
    doc, "• These are non-blocking and run asynchronously in the background", style='List Bullet')

add_paragraph(doc, "C) Callback Queue(s) + Event Loop", style='List Bullet')
add_paragraph(
    doc, "• Holds callbacks once async operations complete", style='List Bullet')
add_paragraph(
    doc, "• Event Loop pushes them to Call Stack when safe", style='List Bullet')

doc.add_paragraph()

# Section 4
add_heading(doc, "4. The 6 Phases of the Event Loop (libuv)")
add_paragraph(doc, "The Event Loop is not a simple loop—it consists of distinct phases, each with its own callback queue. The loop cycles in this fixed order:")

phases = [
    ("1. Timers", "Executes callbacks scheduled by setTimeout() and setInterval()."),
    ("2. Pending Callbacks",
     "Executes I/O callbacks deferred to the next loop iteration (e.g., TCP errors)."),
    ("3. Idle, Prepare", "Used internally by Node.js; not exposed to developers."),
    ("4. Poll", "• Retrieves new I/O events\n• Executes almost all I/O callbacks (e.g., fs, network)\n• May block here waiting for events if no timers are due"),
    ("5. Check", "Executes callbacks queued by setImmediate()."),
    ("6. Close Callbacks", "Executes close event handlers (e.g., socket.on('close', ...)).")
]

for name, desc in phases:
    add_paragraph(doc, f"{name}", style='List Number')
    add_paragraph(doc, desc, style='List Continue')

doc.add_paragraph()

# Section 5
add_heading(doc, "5. Event Loop Flow Diagram (ASCII)")
add_paragraph(doc, "Below is a simplified representation of the loop cycle:")

ascii_diagram = """
   ┌───────────────────────┐
   │        Timers         │  ← setTimeout, setInterval
   └───────────┬───────────┘
               │
   ┌───────────▼───────────┐
   │   Pending Callbacks   │  ← Deferred I/O callbacks
   └───────────┬───────────┘
               │
   ┌───────────▼───────────┐
   │    Idle / Prepare     │  ← Internal
   └───────────┬───────────┘
               │
   ┌───────────▼───────────┐
   │         Poll          │  ← I/O callbacks, may block here
   └───────────┬───────────┘
               │
   ┌───────────▼───────────┐
   │         Check         │  ← setImmediate()
   └───────────┬───────────┘
               │
   ┌───────────▼───────────┐
   │    Close Callbacks    │  ← socket.on('close')
   └───────────┬───────────┘
               │
               └───────► (Next Tick → back to Timers)
"""

add_code_block(doc, ascii_diagram)

doc.add_paragraph(
    "Note: Between each phase, microtasks are processed (see Section 6).")

doc.add_paragraph()

# Section 6
add_heading(doc, "6. Microtasks: nextTick vs Promises")
add_paragraph(doc, "Microtasks have higher priority than regular callbacks and run **after every phase** (and after each callback execution).")

add_paragraph(doc, "Execution Priority Order:")
add_paragraph(doc, "1. process.nextTick() callbacks", style='List Bullet')
add_paragraph(
    doc, "2. Promise.resolve().then() / queueMicrotask()", style='List Bullet')
add_paragraph(doc, "3. Regular event loop phase callbacks",
              style='List Bullet')

add_paragraph(
    doc, "⚠️ Warning: Too many process.nextTick() calls can starve the event loop—avoid recursive nextTick without bounds.")

doc.add_paragraph()

# Section 7
add_heading(doc, "7. setTimeout vs setImmediate – Key Difference")
add_code_block(doc, """// Example
setTimeout(() => console.log('timeout'), 0);
setImmediate(() => console.log('immediate'));

// Output is non-deterministic when run inside main module!
// But inside an I/O callback:
fs.readFile('/file', () => {
  setTimeout(() => console.log('timeout'), 0);
  setImmediate(() => console.log('immediate'));
});
// → Always prints: 'immediate' then 'timeout'
""")

add_paragraph(doc, "Why?")
add_paragraph(
    doc, "• setTimeout(,0) ≈ 1ms delay; runs in Timers phase", style='List Bullet')
add_paragraph(
    doc, "• setImmediate runs in Check phase, which comes **after** Poll", style='List Bullet')
add_paragraph(
    doc, "• When inside an I/O callback (Poll phase), Check runs before next Timers", style='List Bullet')

doc.add_paragraph()

# Section 8
add_heading(doc, "8. Blocking the Event Loop – Common Pitfalls")
add_paragraph(doc, "The Event Loop can be blocked by:")
add_paragraph(
    doc, "• Synchronous APIs: fs.readFileSync, JSON.parse(largeString)", style='List Bullet')
add_paragraph(
    doc, "• CPU-intensive loops: for/while with heavy computation", style='List Bullet')
add_paragraph(
    doc, "• Infinite recursion or unbounded process.nextTick()", style='List Bullet')

add_paragraph(
    doc, "✅ Best Practice: Offload CPU work to Worker Threads or break work into chunks using setImmediate.")

doc.add_paragraph()

# Section 9
add_heading(doc, "9. How the Thread Pool Works")
add_paragraph(doc, "Not all async operations bypass the main thread. libuv maintains a thread pool (default: 4 threads) for operations that don’t support OS-level async, such as:")
add_paragraph(doc, "• fs operations (except fs.realpath.native)",
              style='List Bullet')
add_paragraph(doc, "• DNS lookups (dns.lookup)", style='List Bullet')
add_paragraph(
    doc, "• Crypto operations (crypto.pbkdf2, crypto.randomFill)", style='List Bullet')
add_paragraph(doc, "• zlib compression", style='List Bullet')

add_paragraph(
    doc, "You can adjust pool size via: UV_THREADPOOL_SIZE environment variable (max 128).")

doc.add_paragraph()

# Section 10
add_heading(doc, "10. Practical Examples & Interview Insights")
add_paragraph(doc, "Q: What runs first – Promise.then or setImmediate?")
add_paragraph(
    doc, "A: Promise.then (microtask) runs before setImmediate (Check phase).")

add_paragraph(
    doc, "Q: Can you guarantee execution order between setTimeout and setImmediate?")
add_paragraph(
    doc, "A: Only inside an I/O context. Otherwise, it's timing-dependent.")

add_paragraph(doc, "Q: Is Node.js single-threaded?")
add_paragraph(
    doc, "A: The JavaScript execution is single-threaded, but I/O and thread-pool tasks run in parallel via libuv.")

doc.add_paragraph()

# Section 11
add_heading(doc, "11. Summary – Key Takeaways")
takeaways = [
    "Node.js uses libuv’s event loop to manage async operations.",
    "The loop has 6 well-defined phases; order matters.",
    "Microtasks (nextTick, Promises) always run before moving to the next phase.",
    "Avoid blocking the main thread—use async APIs or Worker Threads.",
    "setImmediate ≠ setTimeout(,0)—phase placement is critical.",
    "Thread pool handles ‘fake async’ operations; not all I/O is truly non-blocking at OS level."
]

for t in takeaways:
    add_paragraph(doc, t, style='List Bullet')

doc.add_paragraph()

# Final note
add_paragraph(doc, "Prepared with technical accuracy for career growth and interview readiness. Target: 12–15 LPA roles.", style='Intense Quote')

# Save document
file_name = "NodeJS_Event_Loop_Notes_by_qwenAI.docx"
doc.save(file_name)
print(f"✅ Document saved as: {file_name}")
print("📁 Check your current working directory for the file.")
